"""
Document Ingestion Module

This script handles ingestion of various document types (PDF, DOCX, TXT, MD),
extracts and cleans their text, splits them into semantically meaningful chunks,
embeds those chunks using a vector model, and stores them in a ChromaDB-backed
persistent vector store for retrieval.

Supported file types:
- .pdf
- .docx
- .txt
- .md

Dependencies:
- PyPDF2, pdfplumber, python-docx, langchain, memory (custom modules)
"""
import os
from PyPDF2 import PdfReader
import pdfplumber
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from memory.embedder import Embedder
from memory.memory_store import MemoryStore
from memory.logger import JSONLogger
import re

embedder = Embedder()
logger = JSONLogger("document_ingestion_log.json")


def extract_pdf_text(file_path):
    """
    Extract text from a PDF file using pdfplumber; fallback to PyPDF2 if necessary.

    Args:
        file_path (str): Path to the PDF file.

    Returns:
        str: Extracted text content.
    """
    try:
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
    except Exception:
        print(f"Failed to extract text from {file_path} using pdfplumber, trying PyPDF2...")
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    
def extract_md_text(file_path):
    """
    Extract raw text from a Markdown (.md) file.

    Args:
        file_path (str): Path to the Markdown file.

    Returns:
        str: Extracted text.
    """
    """Extract text from a Markdown file."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read().strip()

def extract_txt_text(file_path):
    """
    Extract raw text from a plain text (.txt) file.

    Args:
        file_path (str): Path to the text file.

    Returns:
        str: Extracted text.
    """
    """Extract text from a plain text file."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read().strip()

def extract_docx_text(file_path):
    """
    Extract text from a DOCX file including paragraphs and tables.

    Args:
        file_path (str): Path to the DOCX file.

    Returns:
        str: Extracted and formatted text.
    """
    doc = Document(file_path)
    full_text = []

    # Extract text from para with special handling for headings
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading') and para.text.strip():
            full_text.append(f"\n## {para.text.strip()}\n")
        elif para.text.strip():
            full_text.append(para.text.strip())

    # Extract text from each table
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                full_text.append("\t".join(row_text))
    return "\n".join(full_text).strip()

def clean_text(text):
    """
    Clean text by normalizing spaces and line breaks.

    Args:
        text (str): Raw extracted text.

    Returns:
        str: Cleaned text.
    """
    text = re.sub(r'[ \t]+', ' ', text) # Replace multiple spaces/tabs with a single space
    text = re.sub(r'\n{3,}', '\n\n', text) # Replace multiple newlines with two newlines
    return text.strip()


def ingest_file(file_path):
    """
    Ingest a single file:
    - Extracts text
    - Cleans and splits it into chunks
    - Embeds each chunk
    - Stores it in MemoryStore
    - Logs the process

    Args:
        file_path (str): Path to the file to be ingested.

    Raises:
        ValueError: If file type is unsupported.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        text = extract_pdf_text(file_path)
    elif ext == ".md":
        text = extract_md_text(file_path)
    elif ext == ".txt":
        text = extract_txt_text(file_path)
    elif ext == ".docx":
        text = extract_docx_text(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")

    text = clean_text(text)
    
    splitter = RecursiveCharacterTextSplitter( 
        chunk_size=800,
        chunk_overlap=100,
        separators=["\n\n", "\n", ". ", "! ", "? ", ".", "!", "?", " "]
    )

    chunks = splitter.split_text(text)

    total_chunks = len(chunks)
    ingested_chunks = 0

    store = MemoryStore(collection_name="documents")
    for i, chunk in enumerate(chunks):
        if chunk.strip():
            ingested_chunks += 1
            embedding = embedder.get_embedding(chunk)
            metadata = {
                "chunk_index": i
            }
            store.store_document(chunk, embedding, file_path, ext, metadata=metadata)


    logger.log_ingestion(file_path, total_chunks, ingested_chunks)
    print(f"Ingested {ingested_chunks}/{total_chunks} chunks from {file_path}")
